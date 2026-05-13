"""
Generates TradingAlertAssistant_Strategy.docx — complete strategy reference document.
Run once: python generate_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
        # Alternate row shading
        if r_idx % 2 == 0:
            for cell in row_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D6E4F0")
                tcPr.append(shd)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def main():
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("Trading Alert Assistant", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph("Complete Strategy Reference Guide")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    doc.add_paragraph("NSE Intraday | 5-Minute Timeframe | EMA Crossover PRO").runs[0].font.bold = True
    doc.add_paragraph()

    # ── 1. Strategy Overview ─────────────────────────────────────────────────
    add_heading(doc, "1. Strategy Overview", 1, color=(31, 78, 121))
    add_table(doc,
        headers=["Parameter", "Value"],
        rows=[
            ("Strategy Name",   "EMA Crossover PRO"),
            ("Timeframe",        "5-minute candles"),
            ("Fast EMA",         "9"),
            ("Slow EMA",         "21"),
            ("Trend EMA",        "50"),
            ("ATR Period",       "14"),
            ("Risk : Reward",    "1 : 3"),
            ("Volume SMA",       "20"),
        ],
        col_widths=[2.5, 3.5],
    )

    # ── 2. Signal Logic ──────────────────────────────────────────────────────
    add_heading(doc, "2. Signal Logic", 1, color=(31, 78, 121))

    add_heading(doc, "BUY Signal — fires when ALL 3 conditions are true:", 2)
    for point in [
        "EMA 9 crosses ABOVE EMA 21",
        "Price is ABOVE EMA 50 (uptrend confirmed)",
        "Volume > 20-period average (institutional participation)",
    ]:
        p = doc.add_paragraph(point, style="List Bullet")

    add_heading(doc, "SELL Signal — fires when ALL 3 conditions are true:", 2)
    for point in [
        "EMA 9 crosses BELOW EMA 21",
        "Price is BELOW EMA 50 (downtrend confirmed)",
        "Volume > 20-period average",
    ]:
        doc.add_paragraph(point, style="List Bullet")

    # ── 3. Entry, SL & Target ────────────────────────────────────────────────
    add_heading(doc, "3. Entry, Stop Loss & Target", 1, color=(31, 78, 121))
    add_table(doc,
        headers=["Level", "Calculation", "Example (Entry ₹1500, ATR ₹25)"],
        rows=[
            ("Entry",      "Close of signal candle",     "₹1500"),
            ("Stop Loss",  "Entry ± 1× ATR",             "₹1475  (BUY)  |  ₹1525  (SELL)"),
            ("Target",     "Entry ± 3× ATR  (1:3 RR)",   "₹1575  (BUY)  |  ₹1425  (SELL)"),
        ],
        col_widths=[1.5, 2.0, 3.0],
    )

    # ── 4. False Signal Filters ──────────────────────────────────────────────
    add_heading(doc, "4. False Signal Filters (6 Active)", 1, color=(31, 78, 121))
    add_table(doc,
        headers=["#", "Filter", "Value", "Purpose"],
        rows=[
            ("1", "EMA Crossover",   "EMA 9 / 21",      "Entry trigger"),
            ("2", "EMA 50 Trend",    "Above / Below",   "Confirms trend direction"),
            ("3", "Volume Filter",   "> 20 SMA",        "Confirms move strength"),
            ("4", "Min ATR %",       "> 0.3% of price", "Blocks flat / dull stocks"),
            ("5", "ADX Filter",      "> 25",            "Blocks sideways / choppy market"),
            ("6", "Time Filter",     "Before 3:00 PM",  "Blocks end-of-day noise"),
        ],
        col_widths=[0.3, 1.6, 1.6, 2.9],
    )

    doc.add_paragraph(
        "ADX Meaning:  0–20 = Sideways (skip)  |  20–25 = Trend forming  |  "
        "25–40 = Strong trend  |  40+ = Very strong trend"
    ).runs[0].font.italic = True

    # ── 5. Watchlist ─────────────────────────────────────────────────────────
    add_heading(doc, "5. Watchlist — 90 Stocks", 1, color=(31, 78, 121))
    add_table(doc,
        headers=["Sector", "Stocks", "Why Included"],
        rows=[
            ("Nifty 50 Core",     "HDFCBANK, RELIANCE, ICICIBANK, INFY, TCS, BAJFINANCE,\nTATAMOTORS, ADANIENT, HINDALCO, JSWSTEEL, TATASTEEL + more", "Large cap, high liquidity"),
            ("Banking",          "BANKBARODA, PNB, CANBK, FEDERALBNK, IDFCFIRSTB, AUBANK",  "High volume, good ATR"),
            ("IT Midcap",        "PERSISTENT, COFORGE, MPHASIS, LTTS",                       "Strong trending, good ATR"),
            ("Defence",          "HAL, BEL",                                                  "Strong momentum"),
            ("Power",            "ADANIGREEN, TATAPOWER",                                     "High beta, big moves"),
            ("Metals",           "NATIONALUM, HINDZINC",                                      "High ATR %, big intraday range"),
            ("Auto Ancillary",   "MOTHERSON, BALKRISIND",                                     "Clean trends, liquid"),
            ("Capital Goods",    "POLYCAB, DIXON",                                            "High momentum"),
            ("Chemicals",        "DEEPAKNTR, CHOLAFIN",                                       "High ATR, strong trends"),
            ("Realty",           "DLF, GODREJPROP",                                           "Sector momentum"),
            ("Pharma",           "LAURUSLABS, AUROPHARMA, GRANULES",                          "High beta, good range"),
            ("Others",           "ZOMATO, IRCTC, TRENT",                                      "High volume, trendy"),
            ("IT High Momentum", "KPITTECH, TATAELXSI",                                       "Very strong trends, high ATR"),
            ("Capital Goods",    "CGPOWER, ABB, SIEMENS",                                     "Strong trending stocks"),
            ("Defence",          "COCHINSHIP",                                                 "Very high ATR %, big intraday moves"),
            ("Auto",             "ASHOKLEY, TVSMOTOR",                                        "High volume, strong trends"),
            ("Realty",           "PRESTIGE",                                                   "Strong momentum, clean EMA moves"),
            ("Banking",          "RBLBANK",                                                    "High beta, big price moves"),
        ],
        col_widths=[1.5, 3.0, 2.0],
    )

    # ── 6. High Profit Factors ───────────────────────────────────────────────
    add_heading(doc, "6. What Makes This High Profit Focused", 1, color=(31, 78, 121))
    for point in [
        "1:3 Risk Reward — Risk ₹1 to make ₹3. Even a 40% win rate is profitable.",
        "ADX Filter — Only trades when market is genuinely trending. Avoids whipsaws.",
        "High Beta Stocks Added — Metals, Defence, IT Midcap move more = bigger ATR = bigger targets.",
        "Volume Confirmation — Ensures institutional money is behind the move.",
        "No Late Trades — Stops scanning at 3 PM to avoid Zerodha auto squareoff losses.",
        "Min ATR Filter — Skips dull/flat stocks that barely move intraday.",
    ]:
        doc.add_paragraph(point, style="List Bullet")

    # ── 7. Priority Flag System ──────────────────────────────────────────────
    add_heading(doc, "7. Priority Flag System", 1, color=(31, 78, 121))

    doc.add_paragraph(
        "Every Telegram alert is automatically classified as HIGH PRIORITY or Normal "
        "based on whether the stock is in the Top 20 High ATR list. "
        "High ATR stocks move more per candle — giving bigger targets with the same 1:3 RR strategy."
    )
    doc.add_paragraph()

    add_table(doc,
        headers=["Type", "Label in Alert", "What it Means", "Action"],
        rows=[
            ("HIGH PRIORITY", "⭐ HIGH PRIORITY — High ATR Stock", "Top 20 high ATR stock — big moves expected", "Act immediately, higher profit potential"),
            ("Normal",        "(no label)",                        "Good signal but smaller ATR stock",           "Trade if no high priority signal active"),
        ],
        col_widths=[1.3, 2.5, 2.0, 1.7],
    )

    add_heading(doc, "Top 20 High ATR Stocks (Priority List)", 2)
    add_table(doc,
        headers=["Rank", "Stock", "Sector", "Why High Priority"],
        rows=[
            ("1",  "COCHINSHIP",  "Defence",       "Biggest mover, 2–4% daily range"),
            ("2",  "ADANIENT",    "Conglomerate",  "High beta, news-driven"),
            ("3",  "KPITTECH",    "IT",            "Mid-cap momentum, large swings"),
            ("4",  "RBLBANK",     "Banking",       "High beta private bank"),
            ("5",  "TATAMOTORS",  "Auto",          "High volume + range"),
            ("6",  "HINDALCO",    "Metal",         "Commodity-linked volatility"),
            ("7",  "BANKBARODA",  "Banking",       "PSU bank, big moves"),
            ("8",  "ADANIGREEN",  "Power",         "High beta Adani stock"),
            ("9",  "DEEPAKNTR",   "Chemicals",     "High ATR %"),
            ("10", "NATIONALUM",  "Metal",         "High ATR % metal stock"),
            ("11", "JSWSTEEL",    "Metal",         "Commodity volatility"),
            ("12", "TATAPOWER",   "Power",         "Strong momentum"),
            ("13", "IDFCFIRSTB",  "Banking",       "High beta bank"),
            ("14", "LAURUSLABS",  "Pharma",        "High beta pharma"),
            ("15", "CGPOWER",     "Capital Goods", "Strong trending"),
            ("16", "GRANULES",    "Pharma",        "High ATR %"),
            ("17", "TATAELXSI",   "IT",            "Strong trends"),
            ("18", "DLF",         "Realty",        "Sector volatility"),
            ("19", "TATASTEEL",   "Metal",         "Commodity linked"),
            ("20", "ZOMATO",      "Tech",          "High volume, big swings"),
        ],
        col_widths=[0.5, 1.3, 1.5, 3.2],
    )

    # Alert examples
    add_heading(doc, "Alert Examples", 2)

    # HIGH PRIORITY example box
    add_heading(doc, "Example 1 — HIGH PRIORITY Alert (TATAMOTORS)", 3)
    hp_table = doc.add_table(rows=1, cols=1)
    hp_table.style = "Table Grid"
    cell = hp_table.rows[0].cells[0]
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "FFF2CC")
    tcPr.append(shd)
    cell.text = ""
    lines = [
        ("⭐ HIGH PRIORITY  —  High ATR Stock", True,  RGBColor(192, 101, 0)),
        ("🟢 BUY SIGNAL  —  TATAMOTORS (NSE)",  True,  RGBColor(0, 128, 0)),
        ("─────────────────────────────",        False, RGBColor(89, 89, 89)),
        ("📅 Candle  :  16-Apr-2026  10:35 IST  (5-min)", False, None),
        ("📊 ATR     :  ₹14.50",                 False, None),
        ("",                                     False, None),
        ("💰 Entry    :  ₹820.00",               False, None),
        ("🎯 Target   :  ₹863.50  (+5.30%)",     False, None),
        ("🛑 Stop Loss :  ₹805.50  (-1.77%)",    False, None),
        ("",                                     False, None),
        ("⚖️ R : R   :  1 : 3",                  False, None),
        ("─────────────────────────────",        False, RGBColor(89, 89, 89)),
        ("EMA 9/21 cross · EMA 50 trend · Vol filter", False, RGBColor(89, 89, 89)),
    ]
    for i, (text, bold, color) in enumerate(lines):
        if i == 0:
            para = cell.paragraphs[0]
        else:
            para = cell.add_paragraph()
        run = para.add_run(text)
        run.font.bold = bold
        run.font.size = Pt(10)
        if color:
            run.font.color.rgb = color
    doc.add_paragraph()

    # Normal example box
    add_heading(doc, "Example 2 — Normal Alert (INFY)", 3)
    n_table = doc.add_table(rows=1, cols=1)
    n_table.style = "Table Grid"
    cell2 = n_table.rows[0].cells[0]
    tc2 = cell2._tc
    tcPr2 = tc2.get_or_add_tcPr()
    shd2 = OxmlElement("w:shd")
    shd2.set(qn("w:val"), "clear")
    shd2.set(qn("w:color"), "auto")
    shd2.set(qn("w:fill"), "E2EFDA")
    tcPr2.append(shd2)
    cell2.text = ""
    lines2 = [
        ("🟢 BUY SIGNAL  —  INFY (NSE)",         True,  RGBColor(0, 128, 0)),
        ("─────────────────────────────",         False, RGBColor(89, 89, 89)),
        ("📅 Candle  :  16-Apr-2026  11:05 IST  (5-min)", False, None),
        ("📊 ATR     :  ₹6.20",                  False, None),
        ("",                                      False, None),
        ("💰 Entry    :  ₹1482.00",               False, None),
        ("🎯 Target   :  ₹1500.60  (+1.25%)",     False, None),
        ("🛑 Stop Loss :  ₹1475.80  (-0.42%)",    False, None),
        ("",                                      False, None),
        ("⚖️ R : R   :  1 : 3",                   False, None),
        ("─────────────────────────────",         False, RGBColor(89, 89, 89)),
        ("EMA 9/21 cross · EMA 50 trend · Vol filter", False, RGBColor(89, 89, 89)),
    ]
    for i, (text, bold, color) in enumerate(lines2):
        if i == 0:
            para = cell2.paragraphs[0]
        else:
            para = cell2.add_paragraph()
        run = para.add_run(text)
        run.font.bold = bold
        run.font.size = Pt(10)
        if color:
            run.font.color.rgb = color
    doc.add_paragraph()

    # Comparison
    add_heading(doc, "HIGH PRIORITY vs Normal — Key Difference", 2)
    add_table(doc,
        headers=["", "HIGH PRIORITY (TATAMOTORS)", "Normal (INFY)"],
        rows=[
            ("ATR",             "₹14.50",  "₹6.20"),
            ("Stop Loss Risk",  "₹14.50",  "₹6.20"),
            ("Target Profit",   "₹43.50",  "₹18.60"),
            ("% Gain",          "+5.30%",  "+1.25%"),
            ("Action",          "Trade immediately", "Trade if free"),
        ],
        col_widths=[1.5, 2.5, 2.5],
    )

    # ── 8. How to Use Each Alert ─────────────────────────────────────────────
    add_heading(doc, "8. How to Use Each Alert in Zerodha", 1, color=(31, 78, 121))

    add_heading(doc, "When you receive a BUY alert on Telegram:", 2)
    for step in [
        "Open Zerodha Kite → search the stock",
        "Buy at market price immediately (MIS — intraday product)",
        "Go to Positions tab → click Exit on that stock",
        "Change order type to SL-M",
        "Set Trigger Price = Stop Loss value from the alert",
        "Place order — SL fires automatically if price drops",
        "Watch for Target price → exit manually for profit",
    ]:
        doc.add_paragraph(step, style="List Number")

    doc.add_paragraph()
    p = doc.add_paragraph("Important: ")
    p.runs[0].font.bold = True
    p.add_run("Always use MIS product type for intraday. Zerodha auto-squares off all MIS positions at 3:20 PM.")

    # ── 9. Alert System ──────────────────────────────────────────────────────
    add_heading(doc, "9. Alert System", 1, color=(31, 78, 121))
    add_table(doc,
        headers=["Setting", "Value"],
        rows=[
            ("Platform",         "Telegram Bot"),
            ("Market Hours",     "9:15 AM – 3:00 PM IST"),
            ("Scan Frequency",   "Every 6 minutes"),
            ("Duplicate Guard",  "Same signal never sent twice"),
            ("Exchange",         "NSE"),
            ("Data Source",      "TradingView (tvDatafeed)"),
        ],
        col_widths=[2.5, 3.5],
    )

    # ── Footer note ──────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph(
        "Note: This is an alert system, not an auto-trading bot. "
        "All trade executions are manual in Zerodha. "
        "Always use stop loss on every trade."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    # ── Save ─────────────────────────────────────────────────────────────────
    out = "TradingAlertAssistant_Strategy.docx"
    doc.save(out)
    print(f"Document saved: {out}")


if __name__ == "__main__":
    main()
