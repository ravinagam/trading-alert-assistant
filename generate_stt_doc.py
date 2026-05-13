"""
Generates STT_Reduction_Strategies.docx — guide to reduce STT costs in intraday trading.
Run once: python generate_stt_doc.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    if color:
        for run in heading.runs:
            run.font.color.rgb = RGBColor(*color)
    return heading


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(255, 255, 255)
        tc = hdr_cells[i]._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "1F4E79")
        tcPr.append(shd)

    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, cell_text in enumerate(row_data):
            row_cells[c_idx].text = str(cell_text)
        if r_idx % 2 == 0:
            for cell in row_cells:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear")
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:fill"), "D6E4F0")
                tcPr.append(shd)

    if col_widths:
        for i, width in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(width)

    doc.add_paragraph()
    return table


def main():
    doc = Document()

    # ── Title ────────────────────────────────────────────────────────────────
    title = doc.add_heading("STT Reduction Strategies", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)

    sub = doc.add_paragraph("How to Reduce Securities Transaction Tax in Intraday Trading")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.italic = True
    sub.runs[0].font.color.rgb = RGBColor(89, 89, 89)
    doc.add_paragraph()

    # ── 1. What is STT ───────────────────────────────────────────────────────
    add_heading(doc, "1. What is STT?", 1, color=(31, 78, 121))
    p = doc.add_paragraph(
        "STT (Securities Transaction Tax) is a government tax charged on every stock trade in India. "
        "It is automatically deducted by your broker (Zerodha) and cannot be avoided — "
        "but it can be significantly reduced with the right approach."
    )

    doc.add_paragraph()
    add_table(doc,
        headers=["Trade Type", "STT Rate", "Charged On"],
        rows=[
            ("Equity Intraday (MIS)", "0.025%",  "Sell side only"),
            ("Equity Delivery (CNC)", "0.1%",    "Both Buy & Sell"),
            ("Futures",               "0.0125%", "Sell side only"),
            ("Options (Buy)",         "0.0625%", "Premium only"),
        ],
        col_widths=[2.5, 1.5, 2.5],
    )

    # ── 2. How Much STT Are You Paying ───────────────────────────────────────
    add_heading(doc, "2. How Much STT You Could Be Paying", 1, color=(31, 78, 121))
    add_table(doc,
        headers=["Trades/Day", "Avg Trade Value", "STT/Day", "STT/Month (22 days)"],
        rows=[
            ("10 trades", "₹1,00,000", "₹250",  "₹5,500"),
            ("7 trades",  "₹1,00,000", "₹175",  "₹3,850"),
            ("4 trades",  "₹1,00,000", "₹100",  "₹2,200"),
            ("4 trades",  "₹1,50,000", "₹150",  "₹3,300"),
        ],
        col_widths=[1.5, 1.8, 1.5, 2.0],
    )
    doc.add_paragraph(
        "Formula: STT = Trade Value × 0.025% (on sell side for intraday)"
    ).runs[0].font.italic = True

    # ── 3. Strategies to Reduce STT ──────────────────────────────────────────
    add_heading(doc, "3. Strategies to Reduce STT", 1, color=(31, 78, 121))

    # Strategy 1
    add_heading(doc, "Strategy 1: Raise ADX Filter to 25 (Already Done)", 2)
    doc.add_paragraph(
        "ADX > 25 means only strong trending markets generate signals. "
        "This reduces the number of trades per day from many weak signals to fewer, "
        "higher-quality ones — directly cutting STT costs."
    )
    add_table(doc,
        headers=["ADX Setting", "Approx Signals/Day", "STT Impact"],
        rows=[
            ("ADX > 20 (old)", "8–12 signals", "High STT"),
            ("ADX > 25 (new)", "3–5 signals",  "Lower STT, better quality"),
        ],
        col_widths=[2.0, 2.5, 2.0],
    )

    # Strategy 2
    add_heading(doc, "Strategy 2: Trade Fewer but Bigger Positions", 2)
    doc.add_paragraph(
        "Instead of taking every signal with small quantity, pick 3–4 high-conviction "
        "signals per day and trade larger size. Same capital, fewer STT charges."
    )
    for point in [
        "Select only signals where ADX > 30 for larger positions",
        "Skip signals in the first 15 minutes (9:15–9:30) — market is too volatile",
        "Focus on stocks with ATR > 0.5% of price for bigger target potential",
    ]:
        doc.add_paragraph(point, style="List Bullet")

    # Strategy 3
    add_heading(doc, "Strategy 3: Switch to Futures for Large Caps", 2)
    doc.add_paragraph(
        "Futures STT is 0.0125% — exactly half of equity intraday 0.025%. "
        "For large cap stocks, trading the futures contract saves significant STT."
    )
    add_table(doc,
        headers=["Stock", "Equity MIS STT (₹1L)", "Futures STT (₹1L)", "Saving"],
        rows=[
            ("RELIANCE",  "₹25", "₹12.5", "₹12.5 per trade"),
            ("HDFCBANK",  "₹25", "₹12.5", "₹12.5 per trade"),
            ("TATAMOTORS","₹25", "₹12.5", "₹12.5 per trade"),
            ("BANKNIFTY Index Fut", "N/A", "₹12.5", "Best for banking trades"),
        ],
        col_widths=[2.2, 1.8, 1.8, 1.7],
    )
    doc.add_paragraph(
        "Note: Futures require higher margin. Use only if you have sufficient capital."
    ).runs[0].font.italic = True

    # Strategy 4
    add_heading(doc, "Strategy 4: Focus on High ATR Stocks Only", 2)
    doc.add_paragraph(
        "High ATR stocks give bigger price moves per trade. "
        "Profit per trade is larger, making STT a smaller percentage of your gain."
    )
    add_table(doc,
        headers=["Stock", "Why Preferred"],
        rows=[
            ("COCHINSHIP",  "Very high ATR — big intraday moves"),
            ("KPITTECH",    "Strong momentum, large price swings"),
            ("ADANIENT",    "High beta, consistently large ATR"),
            ("TATAMOTORS",  "High volume + high ATR"),
            ("HINDALCO",    "Metal sector volatility"),
            ("BANKBARODA",  "PSU bank, high range"),
            ("RBLBANK",     "High beta banking stock"),
            ("DEEPAKNTR",   "Chemicals — high ATR %"),
        ],
        col_widths=[2.0, 4.5],
    )

    # Strategy 5
    add_heading(doc, "Strategy 5: Avoid Late-Day and Choppy-Day Trading", 2)
    for point in [
        "Already blocked: No signals after 3:00 PM (avoids forced squareoff trades)",
        "On days when NIFTY is flat (range < 100 points) — avoid trading entirely",
        "On expiry days (Thursday) — avoid unless very strong signal, STT adds up fast",
        "Avoid trading in first 15 min (9:15–9:30) — fake breakouts waste capital + STT",
    ]:
        doc.add_paragraph(point, style="List Bullet")

    # ── 4. Summary Comparison ────────────────────────────────────────────────
    add_heading(doc, "4. Before vs After — Monthly STT Estimate", 1, color=(31, 78, 121))
    add_table(doc,
        headers=["Scenario", "Trades/Day", "STT/Month", "Notes"],
        rows=[
            ("Before (ADX > 20)",  "8–10",  "₹4,400–₹5,500", "Many weak signals"),
            ("After  (ADX > 25)",  "3–5",   "₹1,650–₹2,750", "Only strong trends"),
            ("With Futures",       "3–5",   "₹825–₹1,375",   "Half the STT rate"),
        ],
        col_widths=[2.0, 1.5, 2.0, 2.0],
    )
    doc.add_paragraph(
        "Estimated saving: ₹2,000–₹4,000/month just by raising ADX filter to 25."
    ).runs[0].font.bold = True

    # ── 5. Quick Rules ───────────────────────────────────────────────────────
    add_heading(doc, "5. Quick Rules to Follow Daily", 1, color=(31, 78, 121))
    for i, rule in enumerate([
        "Max 4–5 trades per day — quality over quantity",
        "Only take signals where ADX > 25 (strong trend confirmed)",
        "Skip signals on flat NIFTY days (index range < 100 pts)",
        "Skip signals in first 15 min (9:15–9:30 AM)",
        "No trades after 3:00 PM (already enforced by scanner)",
        "For RELIANCE, HDFCBANK, TATAMOTORS — consider futures to halve STT",
        "Target minimum 0.5% profit per trade so STT (0.025%) is less than 5% of gain",
    ], 1):
        doc.add_paragraph(f"{i}. {rule}")

    # ── Footer ───────────────────────────────────────────────────────────────
    doc.add_paragraph()
    note = doc.add_paragraph(
        "STT is a fixed government tax — it cannot be zero. "
        "The goal is to make each trade profitable enough that STT is a small percentage of your gains. "
        "Fewer, stronger trades is always better than many weak trades."
    )
    note.runs[0].font.italic = True
    note.runs[0].font.color.rgb = RGBColor(89, 89, 89)

    out = "STT_Reduction_Strategies.docx"
    doc.save(out)
    print(f"Document saved: {out}")


if __name__ == "__main__":
    main()
